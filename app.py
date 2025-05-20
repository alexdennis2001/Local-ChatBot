import os
import re
from pathlib import Path
from dotenv import load_dotenv

import streamlit as st

import pandas as pd
from docx import Document
from langchain_core.documents import Document as LCDocument

from langchain.agents import AgentExecutor, create_tool_calling_agent
from langchain_core.messages import SystemMessage, AIMessage, HumanMessage
from langchain_core.prompts import PromptTemplate
from langchain_community.vectorstores import SupabaseVectorStore
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain import hub
from langchain_core.tools import tool

from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredExcelLoader,
    UnstructuredWordDocumentLoader,
    UnstructuredCSVLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter

from supabase.client import Client, create_client

# Load environment variables
load_dotenv()

# Initialize Supabase
supabase_url = os.environ.get("SUPABASE_URL")
supabase_key = os.environ.get("SUPABASE_SERVICE_KEY")
supabase: Client = create_client(supabase_url, supabase_key)

# Initialize embeddings and vector store
embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
vector_store = SupabaseVectorStore(
    embedding=embeddings,
    client=supabase,
    table_name="documents",
    query_name="match_documents",
)

# Initialize LLM and prompt
llm = ChatOpenAI(model="gpt-4o", temperature=0)
prompt = hub.pull("hwchase17/openai-functions-agent")

system_message = SystemMessage(
    content=(
        "Eres un ingeniero civil experto. "
        "Respondes siempre en español con un tono profesional pero accesible. "
        "Cuando expliques fórmulas o conceptos matemáticos, SIEMPRE usa formato LaTeX entre doble signo de dólar ($$...$$) o entre \\[ y \\]. "
        "Nunca uses corchetes simples como [ ... ] para fórmulas. "
        "Tu objetivo es ayudar a resolver dudas técnicas, explicar conceptos de ingeniería civil, "
        "y agilizar procesos dentro de la compañía Ingeniería Dennis. "
        "Cuando uses información de documentos, siempre cita la fuente exacta incluyendo el nombre del documento y las líneas aproximadas."
    )
)
prompt.messages.insert(0, system_message)

# Enhanced retrieve tool with line number estimation
@tool(response_format="content_and_artifact")
def retrieve(query: str):
    """Retrieve information related to a query with line number references."""
    retrieved_docs = vector_store.similarity_search(query, k=2)
    
    serialized = []
    for doc in retrieved_docs:
        # Calculate approximate line numbers (assuming ~100 chars per line)
        start_pos = doc.metadata.get('start_index', 0)
        approx_line = start_pos // 100 + 1
        
        source = doc.metadata.get('source', 'desconocida')
        content = doc.page_content.replace('\n', ' ')[:200]  # Single line preview
        serialized.append(
            f"Fuente: {source} (Líneas ~{approx_line}-{approx_line+10})\n"
            f"Contenido: {content}..."
        )
    
    return "\n\n".join(serialized), retrieved_docs

tools = [retrieve]
agent = create_tool_calling_agent(llm, tools, prompt)
agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True, return_intermediate_steps=True)

# Streamlit UI
st.set_page_config(page_title="Ing. Dennis Chatbot", page_icon="🏗️")
st.title("🏗️ Ing. Dennis Chatbot")

# Sidebar uploader with enhanced document processing
st.sidebar.header("📄 Subir documentos")
uploaded_file = st.sidebar.file_uploader(
    "Sube un archivo PDF, TXT, Word, Excel o CSV",
    type=["pdf", "txt", "docx", "xlsx", "csv"]
)

def process_document(file_path: str, file_name: str):
    """Process uploaded document with position tracking"""
    docs = []
    
    try:
        if file_name.endswith(".pdf"):
            loader = PyPDFLoader(file_path)
            raw_docs = loader.load()
            for i, doc in enumerate(raw_docs):
                metadata = {
                    "source": file_name,
                    "page": doc.metadata["page"],
                    "start_index": doc.metadata.get("start_index", i * 1000)
                }
                docs.append(LCDocument(page_content=doc.page_content, metadata=metadata))

        elif file_name.endswith(".txt"):
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            # Split by lines and preserve line numbers
            lines = content.split('\n')
            for i, line in enumerate(lines):
                if line.strip():  # Skip empty lines
                    docs.append(LCDocument(
                        page_content=line,
                        metadata={"source": file_name, "line_number": i+1}
                    ))

        elif file_name.endswith(".docx"):
            doc = Document(file_path)
            full_text = []
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    full_text.append(para.text)
                    docs.append(LCDocument(
                        page_content=para.text,
                        metadata={"source": file_name, "paragraph": i+1}
                    ))
            # Also add full document for context
            docs.append(LCDocument(
                page_content="\n".join(full_text),
                metadata={"source": file_name}
            ))

        elif file_name.endswith(".csv"):
            df = pd.read_csv(file_path)
            for i, row in df.iterrows():
                docs.append(LCDocument(
                    page_content=str(row.to_dict()),
                    metadata={"source": file_name, "row": i+1}
                ))
            # Add full content
            docs.append(LCDocument(
                page_content=df.to_string(index=False),
                metadata={"source": file_name}
            ))

        elif file_name.endswith(".xlsx"):
            df = pd.read_excel(file_path)
            for i, row in df.iterrows():
                docs.append(LCDocument(
                    page_content=str(row.to_dict()),
                    metadata={"source": file_name, "row": i+1}
                ))
            # Add full content
            docs.append(LCDocument(
                page_content=df.to_string(index=False),
                metadata={"source": file_name}
            ))

        # Process chunks with position tracking
        if docs:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=100,
                add_start_index=True
            )
            chunks = splitter.split_documents(docs)
            vector_store.add_documents(chunks)
            return True

    except Exception as e:
        st.sidebar.error(f"❌ Error al procesar el archivo: {str(e)}")
        return False

if uploaded_file:
    os.makedirs("temp_uploads", exist_ok=True)
    file_path = os.path.join("temp_uploads", uploaded_file.name)
    
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    if process_document(file_path, uploaded_file.name):
        st.sidebar.success(f"✅ Documento '{uploaded_file.name}' cargado y procesado con éxito.")

# Chat history
if "messages" not in st.session_state:
    st.session_state.messages = []

for message in st.session_state.messages:
    if isinstance(message, HumanMessage):
        with st.chat_message("user"):
            st.markdown(message.content)
    elif isinstance(message, AIMessage):
        with st.chat_message("assistant"):
            st.markdown(message.content)

# Chat input with enhanced citation handling
if user_question := st.chat_input("Hola soy el ChatBot asistente de Ing. Dennis. ¿En qué puedo ayudarte hoy?"):
    with st.chat_message("user"):
        st.markdown(user_question)
        st.session_state.messages.append(HumanMessage(content=user_question))

    result = agent_executor.invoke({"input": user_question, "chat_history": st.session_state.messages})
    ai_message = result["output"]
    
    # Process citations with line numbers
    citations = []
    for step in result.get("intermediate_steps", []):
        action, observation = step
        
        if isinstance(observation, tuple) and len(observation) == 2:
            retrieved_docs = observation[1]
        else:
            continue
            
        for doc in retrieved_docs:
            if hasattr(doc, "metadata"):
                source = doc.metadata.get("source", "Documento desconocido")
                
                # Get position info based on document type
                if "line_number" in doc.metadata:
                    position = f"Línea {doc.metadata['line_number']}"
                elif "paragraph" in doc.metadata:
                    position = f"Párrafo {doc.metadata['paragraph']}"
                elif "row" in doc.metadata:
                    position = f"Fila {doc.metadata['row']}"
                elif "start_index" in doc.metadata:
                    approx_line = doc.metadata['start_index'] // 100 + 1
                    position = f"Líneas ~{approx_line}-{approx_line+10}"
                else:
                    position = "Posición desconocida"
                
                citation = {
                    'source': source,
                    'position': position,
                    'preview': doc.page_content[:100] + "..."
                }
                if citation not in citations:
                    citations.append(citation)

    with st.chat_message("assistant"):
        # Handle LaTeX formatting
        latex_blocks = re.split(r"(\\\[.*?\\\]|\$\$.*?\$\$)", ai_message, flags=re.DOTALL)
        for block in latex_blocks:
            if block.startswith("\\[") and block.endswith("\\]"):
                st.latex(block[2:-2])
            elif block.startswith("$$") and block.endswith("$$"):
                st.latex(block[2:-2])
            else:
                st.markdown(block)

        # Display enhanced citations
        if citations:
            st.markdown("**📚 Fuentes consultadas:**")
            for cite in citations:
                st.markdown(
                    f"- `{cite['source']}` ({cite['position']})\n"
                    f"  *Extracto:* {cite['preview']}"
                )

    st.session_state.messages.append(AIMessage(content=ai_message))